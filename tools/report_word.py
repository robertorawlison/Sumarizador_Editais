# -*- coding: utf-8 -*-

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt
from io import BytesIO
import subprocess
import ctypes
#import tempfile

def get_ppi_windows():
    hdc = ctypes.windll.user32.GetDC(0)
    ppi = ctypes.windll.gdi32.GetDeviceCaps(hdc, 88)  # 88 é o código para PPI
    ctypes.windll.user32.ReleaseDC(0, hdc)
    return ppi

ppi_windows = get_ppi_windows()

def pixels_to_inches(pixels):
    inches = pixels / ppi_windows
    return inches


def print_word(documents : list) -> None:
    
    # Criando um documento do Word
    docx_document = Document()

    # Adicionando uma tabela ao documento
    table = docx_document.add_table(rows = len(documents) + 1, cols=5)
    table.style = 'TableGrid'

    header_row = table.rows[0]
      
    # Adicionando cabeçalho à tabela
    # Ajustando a largura das colunas
    text_header = ['Doc.', 'Tipo', 'Descrição', 'Data', '#Folhas']
    inche_cells = [Inches(0.25), Inches(1), Inches(4.25), Inches(0.25), Inches(0.25)]
    for cell, text, inche in zip(header_row.cells, text_header, inche_cells):
        cell.width = inche
        cell.text = text
    
    # Aplicando o estilo ao texto nas células do cabeçalho
    for cell in header_row.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Gerando o arquivo temp.doc
    num_lin = 1
    for current_doc in documents:
        row_cells = table.rows[num_lin].cells
        
        # Convertendo a imagem para bytes
        image_bytes = BytesIO()
        current_doc.image.save(image_bytes, format='PNG')

        # Adicionando imagem à célula
        image_stream = BytesIO(image_bytes.getvalue())
        
        image_width, image_height = current_doc.image.size
        inches_width = Inches(pixels_to_inches(2*image_width)/3)
        inches_height = Inches(pixels_to_inches(2*image_height)/3)
        
        #row_cells[0].paragraphs[0].clear()  # Limpar o conteúdo da célula
        row_cells[0].paragraphs[0].add_run().add_picture(image_stream, 
                                                         width = inches_width,
                                                         height = inches_height)

        row_cells[0].add_paragraph(current_doc.get_basename()).alignment = WD_ALIGN_VERTICAL.CENTER
        
        row_cells[1].text = current_doc.type
        row_cells[2].text = current_doc.summary
        row_cells[3].text = current_doc.get_str_date()
        row_cells[4].text = current_doc.get_str_num_pages()
        
        for cell, inche in zip(row_cells, inche_cells):
            cell.width = inche
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    
        
        row_cells[0].paragraphs[1].runs[0].font.size = Pt(5)
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        row_cells[0].height = Inches(pixels_to_inches(image_height + 10)/2)
        
        
        num_lin += 1
        
        
        
    # Salvando o arquivo do Word
    docx_document.save("relatorio_documentos.docx")

    # Abrindo o arquivo gerado
    subprocess.Popen(['start', 'relatorio_documentos.docx'], shell=True)